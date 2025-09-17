from flask import (
    Flask, render_template, render_template_string, request, redirect,
    url_for, send_file, jsonify, session, abort, send_from_directory, make_response,
    flash, current_app
)
from werkzeug.security import generate_password_hash, check_password_hash
from docx import Document
from datetime import timedelta
from dotenv import load_dotenv
from functools import wraps
from jinja2 import TemplateNotFound
from psycopg2.extras import Json
import json
import re
import psycopg2
import pandas as pd
import io
import smtplib
from email.header import Header
from email.utils import formataddr
from email.mime.text import MIMEText
import time, uuid, os
from werkzeug.utils import secure_filename
from pathlib import Path
from uuid import uuid4
from psycopg2.pool import SimpleConnectionPool
import os, psycopg2
from flask import render_template, request, abort
from time import time
import unicodedata
import time

# ========== Flask 应用 ==========
app = Flask(__name__)
try:
    app.json.ensure_ascii = False  # Flask >= 2.3/3.x 推荐写法
except Exception:
    app.config['JSON_AS_ASCII'] = False  # 老版本兜底
app.secret_key = "dev-secret"  # 或者从环境变量读
app.permanent_session_lifetime = timedelta(days=365)

app.config.update(
    SESSION_COOKIE_SAMESITE="Lax",
    SESSION_COOKIE_SECURE=True,   # 仅 HTTPS 时安全发送
)
if os.getenv("FLASK_ENV") != "production":
    app.config["SESSION_COOKIE_SECURE"] = False

# —— 常见“中文→UTF-8→被按latin1读”的假字符标记
_MOJIBAKE_MARKERS = ("Ã", "Â", "â€™", "â€œ", "â€", "å¼", "æ", "ç", "�")

def _maybe_fix_encoding(s: str) -> str:
    """把 'å¼'、'Ã§' 这类假字符尽量还原为真正UTF-8中文，并清理控制符。"""
    if not isinstance(s, str):
        return s
    t = s
    # 1) 只有看起来像乱码才尝试回转
    if any(m in s for m in _MOJIBAKE_MARKERS):
        for enc in ("latin1", "cp1252"):
            try:
                t2 = s.encode(enc, errors="strict").decode("utf-8")
                # 如果回转后不再包含常见乱码标记，就认为成功
                if not any(m in t2 for m in _MOJIBAKE_MARKERS):
                    t = t2
                    break
            except Exception:
                pass
    # 2) 统一做Unicode规范化，并去掉不可见控制字符
    try:
        t = unicodedata.normalize("NFC", t)
    except Exception:
        pass
    # 清掉不可见控制符（保留换行\t）
    import re as _re
    t = _re.sub(r"[\u0000-\u0008\u000b\u000c\u000e-\u001f\u007f]", "", t)
    return t

def _normalize_obj(x):
    """递归归一化：dict/list/tuple/str 都处理；其他类型原样返回。"""
    if x is None or isinstance(x, (int, float, bool)):
        return x
    if isinstance(x, str):
        return _maybe_fix_encoding(x)
    if isinstance(x, dict):
        return { _normalize_obj(k): _normalize_obj(v) for k, v in x.items() }
    if isinstance(x, (list, tuple, set)):
        return [ _normalize_obj(v) for v in x ]
    return x


# 放在 app 初始化后
@app.context_processor
def inject_asset():
    import os
    def asset(path):
        full = os.path.join(app.static_folder, path)
        v = str(int(os.path.getmtime(full))) if os.path.exists(full) else str(int(time()))
        return url_for("static", filename=path, v=v)
    return {"asset": asset}

UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["MAX_CONTENT_LENGTH"] = 50 * 1024 * 1024  # 50MB

# ========== 基本配置 ==========
load_dotenv()
app.secret_key = os.getenv("SECRET_KEY", "replace-this-in-prod")

# 连接池：复用到 Neon 的连接，避免每次请求建链
DB_URL = os.getenv("DB_URL", globals().get("DB_URL", ""))
if not DB_URL:
    raise RuntimeError("DB_URL is not set. Please configure your database DSN.")
_POOL = SimpleConnectionPool(minconn=1, maxconn=int(os.getenv("DB_POOL_MAX","10")), dsn=DB_URL)


class _ConnProxy:
    __slots__ = ("_raw",)
    def __init__(self, raw):
        self._raw = raw
    def __getattr__(self, name):
        return getattr(self._raw, name)
    def close(self):
        # 把 close 变成“归还连接到连接池”
        try:
            _POOL.putconn(self._raw)
        except Exception:
            try:
                self._raw.close()
            except Exception:
                pass

def get_conn():
    conn = _POOL.getconn()
    # 会话级安全优化（失败忽略）
    try:
        conn.set_client_encoding('UTF8')
    except Exception:
        pass
    try:
        with conn.cursor() as c:
            c.execute("SET statement_timeout TO 60000")                     # 60s
            c.execute("SET idle_in_transaction_session_timeout TO 30000")   # 30s
    except Exception:
        pass
    return _ConnProxy(conn)

# Gzip 压缩与静态缓存（不改业务逻辑）
try:
    from flask_compress import Compress
    Compress(app)
    app.config["COMPRESS_MIMETYPES"] = [
        "text/html", "text/css", "application/json",
        "application/javascript", "image/svg+xml"
    ]
    app.config["COMPRESS_LEVEL"] = 6
    app.config["COMPRESS_MIN_SIZE"] = 1024
except Exception as _e:
    print("[perf] flask-compress not active:", _e)

try:
    from datetime import timedelta
    app.config["SEND_FILE_MAX_AGE_DEFAULT"] = timedelta(days=7)
except Exception:
    pass

@app.after_request
def _perf_add_cache_headers(resp):
    try:
        ct = resp.headers.get("Content-Type", "")
        # HTML 不缓存（避免管理端看到旧页）
        if "text/html" in ct:
            resp.headers.setdefault("Cache-Control", "no-store")
        # 静态资源缓存一周（第一次加载后明显加速）
        elif "javascript" in ct or "css" in ct or "image/" in ct or "svg+xml" in ct or "font/" in ct:
            resp.headers.setdefault("Cache-Control", "public, max-age=604800, immutable")
    except Exception:
        pass
    return resp

# === UTF-8 兜底(1)：强制所有 HTML/JSON 响应带 charset=utf-8 ===
@app.after_request
def _force_utf8_headers(resp):
    try:
        ct = resp.headers.get("Content-Type", "")
        low = ct.lower()
        if (low.startswith("text/") or "application/json" in low) and "charset=" not in low:
            base = ct.split(";")[0].strip() or "text/plain"
            resp.headers["Content-Type"] = f"{base}; charset=utf-8"
    except Exception:
        pass
    return resp

# === UTF-8 兜底(2)：HTML 没有 <meta charset> 时自动注入 ===
@app.after_request
def _ensure_meta_charset(resp):
    try:
        if resp.content_type and resp.content_type.startswith("text/html"):
            html = resp.get_data(as_text=True)
            low = html.lower()
            if "</head>" in low and "<meta charset=" not in low:
                html = html.replace("</head>", '<meta charset="utf-8"></head>', 1)
                resp.set_data(html)
    except Exception:
        pass
    return resp

# SMTP（可选）
SMTP_SERVER = os.getenv("SMTP_SERVER", "")
SMTP_PORT = int(os.getenv("SMTP_PORT", "587") or 587)
SENDER_EMAIL = os.getenv("SENDER_EMAIL", "")
SENDER_PASSWORD = os.getenv("SENDER_PASSWORD", "")

# ========== 用户表 ==========
def init_user_table():
    conn = get_conn(); c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS users (
        id SERIAL PRIMARY KEY,
        username TEXT UNIQUE,
        password_hash TEXT,
        role TEXT DEFAULT 'admin'
    )''')
    conn.commit(); conn.close()
init_user_table()

# ========= 主 submissions（保留）=========
def init_main_submissions():
    conn = get_conn(); c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS submissions (
        id SERIAL PRIMARY KEY,
        name TEXT, phone TEXT, email TEXT,
        group_name TEXT, event_name TEXT,
        start_date TEXT, start_time TEXT, end_date TEXT, end_time TEXT,
        location TEXT, event_type TEXT, participants TEXT,
        equipment TEXT, special_request TEXT,
        donation TEXT, donation_method TEXT, remarks TEXT,
        emergency_name TEXT, emergency_phone TEXT, attachment TEXT,
        status TEXT DEFAULT '待审核',
        review_comment TEXT
    )''')
    conn.commit(); conn.close()
init_main_submissions()

# ========= form_defs =========
def init_form_defs():
    conn = get_conn(); c = conn.cursor()
    # 统一表结构：schema_json 用 JSONB
    c.execute("""
        CREATE TABLE IF NOT EXISTS form_defs (
            id          SERIAL PRIMARY KEY,
            name        TEXT        NOT NULL,
            site_name   TEXT UNIQUE NOT NULL,
            schema_json JSONB       NOT NULL DEFAULT '{}'::jsonb,
            created_by  INT REFERENCES users(id),
            db_url      TEXT NOT NULL,
            description TEXT,
            created_at  TIMESTAMP   DEFAULT CURRENT_TIMESTAMP
        )
    """)
    # 如果历史库里 schema_json 还是 TEXT，这里在线迁移到 JSONB
    c.execute("""
        SELECT data_type
          FROM information_schema.columns
         WHERE table_name='form_defs' AND column_name='schema_json'
    """)
    row = c.fetchone()
    if row and row[0].lower() in ("text", "character varying"):
        c.execute("""
            ALTER TABLE form_defs
            ALTER COLUMN schema_json TYPE JSONB
            USING CASE
                    WHEN schema_json IS NULL OR schema_json='' THEN '{}'::jsonb
                    ELSE schema_json::jsonb
                 END
        """)
    # 兜底补列
    c.execute("ALTER TABLE form_defs ADD COLUMN IF NOT EXISTS description TEXT")
    c.execute("ALTER TABLE form_defs ADD COLUMN IF NOT EXISTS created_at  TIMESTAMP DEFAULT CURRENT_TIMESTAMP")
    c.execute("ALTER TABLE form_defs ADD COLUMN IF NOT EXISTS created_by  INT")
    c.execute("ALTER TABLE form_defs ADD COLUMN IF NOT EXISTS db_url      TEXT")
    conn.commit(); conn.close()
init_form_defs()

# ========== 权限 ==========
def admin_required(view_func):
    @wraps(view_func)
    def wrapper(*args, **kwargs):
        if not session.get("user_id"):
            nxt = request.path if request.method == "GET" else "/index"
            return redirect(url_for("login", enter=1, next=nxt))
        return view_func(*args, **kwargs)
    return wrapper

# === 变更点 ③：全局登录门禁（除公共页外统统要求登录） ===
@app.before_request
def _global_login_gate():
    """
    放行公共访问的 URL：
      - /login, /register
      - /f/<site_name> 公开表单 GET/POST
      - /site/<site_name>/status_query 公开状态查询
      - /site/<site_name>/uploads/* 公开文件访问
      - /site/<site_name>/draft/save 公开草稿保存（前台填写用）
      - /uploads/* （历史数据/兼容）
      - /static/*, /favicon.ico, /robots.txt
      - /_health 健康检查
    其它未登录访问将 302 跳转到 /login?enter=1&next=...
    """
    path = (request.path or "/")
    # 快速白名单
    if (
        path.startswith("/static/")
        or path in ("/favicon.ico", "/robots.txt", "/_health")
        or path.startswith("/login")
        or path.startswith("/register")
        or path.startswith("/uploads/")
        or path.startswith("/f/")
    ):
        return None

    # /site/<site_name>/... 的细粒度白名单
    m = re.match(r"^/site/([^/]+)/(.*)$", path)
    if m:
        tail = m.group(2)
        if (
            tail.startswith("status_query")
            or tail.startswith("uploads/")
            or tail == "draft/save"
        ):
            return None  # 公共放行

    # 未登录统一拦截
    if not session.get("user_id"):
        # 避免死循环
        if not path.startswith("/login"):
            nxt = request.full_path if request.query_string else path
            # 安全 next 仅保留站内路径
            try:
                if not isinstance(nxt, str) or not nxt.startswith("/"):
                    nxt = "/index"
            except Exception:
                nxt = "/index"
            return redirect(url_for("login", enter=1, next=nxt))
    return None
# === 全局登录门禁 END ===

# ---------- 开场动画 ----------
SPLASH_GATE_HTML = r"""
<!doctype html>
<meta charset="utf-8">
<title>Formly</title>
<style>
  :root{
    --splash-appear: 1500ms;
    --splash-hold:   1600ms;
    --splash-exit:   1100ms;
    --splash-total: calc(var(--splash-appear) + var(--splash-hold) + var(--splash-exit));
  }
  html,body{height:100%}
  body{
    margin:0;
    font-family: "Inter","Microsoft YaHei", ui-sans-serif, -apple-system, system-ui, "Segoe UI", Roboto, Helvetica, Arial;
    background:
      radial-gradient(900px 600px at -10% -10%, rgba(255, 214, 231, .28), transparent 60%),
      radial-gradient(900px 600px at 110% -10%, rgba(214, 241, 255, .26), transparent 60%),
      radial-gradient(900px 600px at 50% 120%, rgba(253, 233, 217, .26), transparent 60%),
      linear-gradient(180deg, #ffffff 0%, #fffcfa 45%, #f7fbff 100%);
    display:flex; align-items:flex-start; justify-content:center; overflow:hidden;
  }
  #splash{ position:relative; width:100%; text-align:center; top:120vh; animation: slide var(--splash-total) forwards; }
  #splash h1{
    margin:0; padding:0 20px; font-size: clamp(40px, 9vw, 82px); font-weight:1000; letter-spacing:.6px; line-height:1.08;
    color:transparent; background: linear-gradient(90deg, #ff9aa2 0%, #fecfef 35%, #f6d365 70%, #fda085 100%);
    -webkit-background-clip:text; background-clip:text; text-shadow:0 6px 22px rgba(255,154,162,.28);
  }
  #splash p{ margin:14px 0 0; color:#5b6472; font-weight:700; font-size:clamp(12px,2.6vw,16px); }
  @keyframes slide{
    0%{top:120vh; filter:blur(10px); opacity:0;}
    38%{top:42vh; filter:blur(4px); opacity:1; animation-timing-function:cubic-bezier(.2,1,.2,1);}
    58%{top:40vh; filter:blur(0); opacity:1;}
    70%{top:40vh;}
    100%{top:-60vh; filter:blur(2px); opacity:0; animation-timing-function:cubic-bezier(.4,0,1,1);}
  }
</style>
<div id="splash">
  <h1>Formly平台</h1>
  <p>设计属于您的表单</p>
</div>
<script>
  const LOGIN = "{{ login_url }}";
  const total = get('--splash-total', 4200);
  setTimeout(()=>location.replace(LOGIN), total + 200);
  function get(name, d){ const v=getComputedStyle(document.documentElement).getPropertyValue(name).trim();
    if(!v) return d; if(v.endsWith('ms')) return parseFloat(v); if(v.endsWith('s')) return parseFloat(v)*1000;
    const n=parseFloat(v); return isNaN(n)?d:n;
  }
</script>
"""

# ========= 首页：只在第一次显示开场 =========
@app.route("/")
def landing():
    if session.get("user_id"):
        return redirect(url_for("index"))
    if request.cookies.get("saw_splash") == "1":
        return redirect(url_for("login", enter=1))
    resp = make_response(render_template_string(
        SPLASH_GATE_HTML, login_url=url_for("login", enter=1)
    ))
    resp.set_cookie("saw_splash", "1", max_age=60*60*24*90, samesite="Lax")
    return resp

# ========= 登录 / 注册 =========
@app.route("/login", methods=["GET", "POST"])
def login():
    error = None
    if request.method == "POST":
        username = request.form.get("username")
        password = request.form.get("password")

        conn = get_conn(); c = conn.cursor()
        c.execute("SELECT id, password_hash, role FROM users WHERE username=%s", (username,))
        row = c.fetchone(); conn.close()

        if row and check_password_hash(row[1], password):
            session.permanent = True
            session["user_id"] = row[0]
            session["username"] = username
            session["role"] = row[2] or "admin"
            next_url = request.args.get("next") or request.form.get("next") or "/index"
            if next_url == "/" or not (isinstance(next_url, str) and next_url.startswith("/")):
                next_url = "/index"
            sep = "&" if "?" in next_url else "?"
            return redirect(f"{next_url}{sep}enter=1")
        else:
            error = "❌ 用户名或密码错误"

    try:
        return render_template("login.html", error=error)
    except TemplateNotFound:
        return render_template("index.html", error=error)

@app.route("/register", methods=["GET", "POST"])
def register():
    next_url = request.args.get("next") or request.form.get("next") or "/index"
    if request.method == "GET":
        return redirect(url_for("login", enter=1, next=next_url))

    username = (request.form.get("username") or "").strip()
    password = (request.form.get("password") or "").strip()

    if not username or not password:
        flash("❌ 用户名和密码不能为空")
        return redirect(url_for("login", enter=1, next=next_url))

    conn = get_conn(); c = conn.cursor()
    try:
        c.execute("SELECT id FROM users WHERE username=%s", (username,))
        if c.fetchone():
            flash("❌ 用户已存在")
            conn.close()
            return redirect(url_for("login", enter=1, next=next_url))

        c.execute(
            "INSERT INTO users (username, password_hash, role) VALUES (%s,%s,%s) RETURNING id",
            (username, generate_password_hash(password), "admin")
        )
        conn.commit()
    except Exception as e:
        conn.rollback()
        flash(f"❌ 注册失败：{e}")
        return redirect(url_for("login", enter=1, next=next_url))
    finally:
        conn.close()

    flash("✅ 注册成功，请登录")
    return redirect(url_for("login", enter=1, next=next_url))

@app.route("/logout")
def logout():
    session.clear()
    resp = make_response(redirect(url_for("landing")))
    resp.delete_cookie("saw_splash")
    return resp

# ========= 管理首页 =========
@app.route("/index")
@admin_required
def index():
    role = session.get("role")
    forms = []
    if role == "admin":
        user_id = session.get("user_id")
        conn = get_conn(); c = conn.cursor()
        c.execute("""
            SELECT id, name, site_name, db_url, created_at
            FROM form_defs
            WHERE created_by=%s
            ORDER BY id ASC
        """, (user_id,))
        rows = c.fetchall(); conn.close()
        forms = [{
            "id": r[0], "name": r[1], "site_name": r[2], "db_url": r[3],
            "created_at": str(r[4]) if r[4] else "-"
        } for r in rows]
    return render_template("dashboard.html",
                           role=role,
                           username=session.get("username"),
                           forms=forms)

# ========= 入场动画注入 =========
@app.after_request
def _inject_enter_animation(resp):
    try:
        path = (request.path or "")
        if not resp.content_type.startswith("text/html"):
            return resp
        need = (path == "/index" or path == "/create_form" or (path.startswith("/site/") and path.endswith("/admin")))
        if not need:
            return resp

        html = resp.get_data(as_text=True)
        inject = """
        <style id="page-enter-style">
        @keyframes riseSoft{
          0%   { transform: translateY(40px); opacity:0; filter: blur(6px); }
          100% { transform: translateY(0);    opacity:1; filter: blur(0); }
        }
        @keyframes riseAndSettle{
          0%   { transform: translateY(60vh) scale(.98); opacity:0; filter: blur(10px); }
          55%  { transform: translateY(-1.5vh) scale(1); opacity:1; filter: blur(0);
                 animation-timing-function:cubic-bezier(.2,1,.2,1); }
          72%  { transform: translateY(.6vh); }
          100% { transform: translateY(0); }
        }
        html.page-enter, body.page-enter{ overflow-x:hidden; }
        </style>
        <script>
        (function(){
          var p = new URLSearchParams(location.search);
          var should = p.get('enter') === '1';
          try {
            var ref = document.referrer && new URL(document.referrer).pathname;
            if (!should && ref === '/index') should = true;
          } catch(e){}
          if(!should) return;
          document.documentElement.classList.add('page-enter');
          document.body.classList.add('page-enter');
          function go(){
            var main = document.querySelector(
              'main, .card, .box, .container, .wrap, .page, #app, #root, body > div, body > main'
            ) || document.body.firstElementChild;
            if(main){
              main.style.willChange = 'transform,opacity,filter';
              var anim = (location.pathname === '/index')
                ? 'riseSoft 1600ms cubic-bezier(.22,.95,.24,1) forwards'
                : 'riseAndSettle 1600ms forwards';
              main.style.animation = anim;
            }
            if (p.get('enter') === '1') {
              try{ history.replaceState(null,'', location.pathname); }catch(e){}
            }
          }
          if(document.readyState==='loading'){ document.addEventListener('DOMContentLoaded', go); } else { go(); }
        })();
        </script>
        """
        if "</body>" in html:
            resp.set_data(html.replace("</body>", inject + "</body>"))
    except Exception:
        pass
    return resp

# ========= 后台主题色注入 & 去除块 =========
@app.after_request
def _inject_admin_theme_and_cleanup(resp):
    try:
        path = (request.path or "")
        if not resp.content_type.startswith("text/html"):
            return resp

        m = re.match(r"^/site/([^/]+)/admin/?$", path)
        if not m:
            return resp
        site_name = m.group(1)

        brand = "#2563eb"
        try:
            conn = get_conn(); c = conn.cursor()
            c.execute("SELECT schema_json FROM form_defs WHERE site_name=%s", (site_name,))
            row = c.fetchone(); conn.close()
            if row:
                schema = row[0] if isinstance(row[0], dict) else (json.loads(row[0]) if row[0] else {})
                b = (((schema or {}).get("theme") or {}).get("brand") or "").strip()
                if b:
                    brand = b
        except Exception:
            pass

        html = resp.get_data(as_text=True)
        inject = f"""
<style id="injected-theme-brand">
  :root{{ --accent: {brand}; }}
  body{{ background: {brand} !important; }}
</style>
<script>
(function(){{
  const killTitles = ['背景样式','通知'];
  function removeBlocks(){{
    const heads = document.querySelectorAll('h1,h2,h3,h4,.title,.card-title');
    heads.forEach(h => {{
      const text = (h.textContent||'').trim();
      if(killTitles.some(k => text.includes(k))) {{
        const card = h.closest('.card, section, article, .box, .panel, .container, .wrap, div');
        if (card && card.parentNode) card.parentNode.removeChild(card);
      }}
    }});
  }}
  if(document.readyState==='loading') document.addEventListener('DOMContentLoaded', removeBlocks);
  else removeBlocks();
}})();
</script>
"""
        if "</body>" in html:
            resp.set_data(html.replace("</body>", inject + "</body>"))
    except Exception:
        pass
    return resp

# ========== 安全 next ==========
def _safe_next_path(path: str) -> str:
    if path and isinstance(path, str) and path.startswith("/"):
        return path
    return url_for("index")

@app.after_request
def _inject_preview_guard(resp):
    try:
        if request.path.endswith("/preview") and resp.content_type.startswith("text/html"):
            html = resp.get_data(as_text=True)
            inject = r"""
<style id="preview-guard-style">
  #preview-banner{
    position:fixed;left:50%;top:14px;transform:translateX(-50%);
    background:#111827;color:#fff;border-radius:999px;padding:8px 14px;
    font-weight:800;font-size:13px;z-index:9999;box-shadow:0 6px 20px rgba(0,0,0,.2);
    opacity:.92
  }
</style>"""
            if "</body>" in html:
                resp.set_data(html.replace("</body>", inject + "</body>"))
    except Exception:
        pass
    return resp


# ========== 创建/编辑 ==========
@app.route("/create_form", methods=["GET", "POST"])
@admin_required
def create_form():
    user_id = session.get("user_id")

    # === GET: 成功页 ===
    if request.method == "GET" and request.args.get("saved") == "1":
        site = (request.args.get("site") or "").strip()
        if not site:
            return redirect(url_for("create_form_new", enter=1))
        public_url = url_for("public_form", site_name=site, _external=True)
        admin_url  = url_for("create_form", site=site, _external=True)
        return render_template("create_success.html",
                               site_name=site, public_url=public_url, admin_url=admin_url)

    # === GET: new=1 ===
    if request.method == "GET" and request.args.get("new") == "1":
        return render_template(
            "create_form.html",
            form_name=None, form_desc=None, site_name=None,
            schema_json=None, submissions=[],
        )

    # === POST: 创建 / 更新 ===
    if request.method == "POST":
        is_ajax = (
            request.args.get("ajax") == "1" or
            request.headers.get("X-Requested-With") == "XMLHttpRequest" or
            "application/json" in (request.headers.get("Accept") or "")
        )

        name        = (request.form.get("form_name") or "").strip()
        site_name   = (request.form.get("site_name") or "").strip()
        form_desc   = (request.form.get("form_desc") or "").strip()
        schema_json = request.form.get("schema_json") or '{"fields": []}'

        if not name or not site_name:
            msg = "❌ 表单名称或网站名不能为空"
            return (jsonify({"ok": False, "error": msg}), 400) if is_ajax else (msg, 400)

        if not re.fullmatch(r"[A-Za-z_][A-Za-z0-9_]*", site_name):
            msg = "❌ 站点名只能包含字母、数字与下划线，且不能以数字开头"
            return (jsonify({"ok": False, "error": msg}), 400) if is_ajax else (msg, 400)

        schema_name = _safe_schema(site_name)
        try:
            schema_obj = json.loads(schema_json) if isinstance(schema_json, str) else (schema_json or {})
        except Exception:
            schema_obj = {}

        conn = get_conn(); c = conn.cursor()
        try:
            # 确保表存在 & 兼容
            c.execute("""
                CREATE TABLE IF NOT EXISTS form_defs (
                    id          SERIAL PRIMARY KEY,
                    name        TEXT        NOT NULL,
                    site_name   TEXT UNIQUE NOT NULL,
                    schema_json JSONB       NOT NULL DEFAULT '{}'::jsonb,
                    created_by  INT,
                    db_url      TEXT,
                    description TEXT,
                    created_at  TIMESTAMP   DEFAULT CURRENT_TIMESTAMP
                )
            """)
            c.execute("ALTER TABLE form_defs ADD COLUMN IF NOT EXISTS created_by  INT")
            c.execute("ALTER TABLE form_defs ADD COLUMN IF NOT EXISTS db_url      TEXT")
            c.execute("ALTER TABLE form_defs ADD COLUMN IF NOT EXISTS description TEXT")
            c.execute("ALTER TABLE form_defs ADD COLUMN IF NOT EXISTS created_at  TIMESTAMP DEFAULT CURRENT_TIMESTAMP")
            c.execute("ALTER TABLE form_defs ADD COLUMN IF NOT EXISTS schema_json JSONB DEFAULT '{}'::jsonb")
            c.execute("""
                SELECT data_type FROM information_schema.columns
                 WHERE table_name='form_defs' AND column_name='schema_json'
            """)
            t = c.fetchone()
            if t and t[0].lower() in ("text", "character varying"):
                c.execute("""
                    ALTER TABLE form_defs
                    ALTER COLUMN schema_json TYPE JSONB
                    USING CASE
                            WHEN schema_json IS NULL OR schema_json='' THEN '{}'::jsonb
                            ELSE schema_json::jsonb
                         END
                """)

            # UPSERT
            c.execute("""
                INSERT INTO form_defs (name, site_name, schema_json, created_by, db_url, description)
                VALUES (%s, %s, %s, %s, %s, %s)
                ON CONFLICT (site_name) DO UPDATE SET
                    name        = EXCLUDED.name,
                    schema_json = EXCLUDED.schema_json,
                    created_by  = EXCLUDED.created_by,
                    db_url      = EXCLUDED.db_url,
                    description = EXCLUDED.description
                RETURNING id
            """, (name, site_name, Json(schema_obj), user_id, schema_name, form_desc))
            _ = c.fetchone()[0]

            # 业务 schema + submissions
            c.execute(f"CREATE SCHEMA IF NOT EXISTS {schema_name}")
            c.execute(f"""
                CREATE TABLE IF NOT EXISTS {schema_name}.submissions (
                    id             SERIAL PRIMARY KEY,
                    user_id        INT,
                    data           JSONB,
                    status         TEXT DEFAULT '待审核',
                    review_comment TEXT,
                    created_at     TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)
            # 草稿表（供 /draft/save）
            c.execute(f"""
                CREATE TABLE IF NOT EXISTS {schema_name}.drafts (
                    token TEXT PRIMARY KEY,
                    data  JSONB,
                    files JSONB,
                    created_at TIMESTAMP DEFAULT NOW(),
                    updated_at TIMESTAMP DEFAULT NOW()
                )
            """)
            conn.commit()
        except Exception as e:
            conn.rollback()
            try: current_app.logger.exception("create_form failed")
            except Exception: pass
            return (jsonify({"ok": False, "error": f"❌ 创建失败: {e}"}), 500) if is_ajax else (f"❌ 创建失败: {e}", 500)
        finally:
            conn.close()

        public_url = url_for("public_form", site_name=site_name, _external=True)
        admin_url  = url_for("create_form", site=site_name, _external=True)

        wants_json = (
            request.args.get("ajax") == "1" or
            request.headers.get("X-Requested-With") == "XMLHttpRequest" or
            "application/json" in (request.headers.get("Accept") or "").lower()
        )

        if wants_json:
            return jsonify({
                "ok": True,
                "site_name": site_name,
                "public_url": public_url,
                "admin_url": admin_url,
                "success_url": url_for("create_success", site_name=site_name)
            })

        return render_template(
            "create_success.html",
            site_name=site_name,
            public_url=public_url,
            admin_url=admin_url,
        )

    # 普通 GET：按 site 或最近一份
    site_q = (request.args.get("site") or "").strip()
    form_name = form_desc = site_name = schema_json = None
    submissions = []

    conn = get_conn(); c = conn.cursor()
    try:
        if site_q:
            c.execute("""
                SELECT name, site_name, schema_json, COALESCE(description,''::TEXT)
                  FROM form_defs
                 WHERE site_name=%s
                 LIMIT 1
            """, (site_q,))
            row = c.fetchone()
            if not row:
                return redirect(url_for("create_form_new", enter=1))
            form_name, site_name, schema_json, form_desc = row
        else:
            c.execute("""
                SELECT name, site_name, schema_json, COALESCE(description,''::TEXT)
                  FROM form_defs
                 WHERE created_by=%s
                 ORDER BY id DESC
                 LIMIT 1
            """, (user_id,))
            row = c.fetchone()
            if not row:
                return redirect(url_for("create_form_new", enter=1))
            form_name, site_name, schema_json, form_desc = row

        try:
            schema_name = _safe_schema(site_name)
            c.execute(f"SET search_path TO {schema_name}")
            c.execute("SELECT id, user_id, data, status, created_at FROM submissions ORDER BY id DESC LIMIT 50")
            submissions = c.fetchall()
        except Exception:
            pass
    finally:
        conn.close()

    return render_template(
        "create_form.html",
        form_name=form_name, form_desc=form_desc,
        site_name=site_name, schema_json=schema_json,
        submissions=submissions,
    )

TEMPLATES = {
    "blank":   {"form_name":"空白表单","form_desc":"进入后可自由添加/删除题目。","schema":{"fields":[]}},
    "contact": {"form_name":"联系信息","form_desc":"请留下您的联系方式，我们会尽快回复。","schema":{"fields":[
        {"key":"name","label":"姓名","type":"text","required":True},
        {"key":"phone","label":"手机","type":"tel"},
        {"key":"email","label":"邮箱","type":"email","required":True},
        {"key":"msg","label":"留言","type":"textarea","required":True},
    ]}},
    "feedback":{"form_name":"问题反馈","form_desc":"感谢反馈！请尽量详细描述问题。","schema":{"fields":[
        {"key":"type","label":"类型","type":"select","options":["功能建议","体验问题","Bug","其他"],"required":True},
        {"key":"desc","label":"详细描述","type":"textarea","required":True},
        {"key":"shots","label":"截图/视频","type":"file","accept":"jpg,png,webp,mp4,webm","maxFiles":5},
    ]}},
    "event":   {"form_name":"活动报名","form_desc":"请填写报名信息并同意条款。","schema":{"fields":[
        {"key":"name","label":"姓名","type":"text","required":True},
        {"key":"email","label":"邮箱","type":"email","required":True},
        {"key":"phone","label":"手机","type":"tel","required":True},
        {"key":"ticket","label":"票种","type":"select","options":["早鸟票","标准票","VIP票"],"required":True},
    ]}},
}

@app.get("/create_form/new")
@admin_required
def create_form_new():
    tpl = request.args.get("tpl")
    if tpl:
        t = TEMPLATES.get(tpl)
        if t:
            return render_template(
                "create_form.html",
                form_name=t["form_name"],
                form_desc=t["form_desc"],
                schema_json=t["schema"],
                site_name="",
                tpl=tpl
            )
    # 无 tpl -> 空白
    return render_template(
        "create_form.html",
        form_name=None, form_desc=None, site_name=None,
        schema_json=None, submissions=[],
    )

# Postgres 标识符最长 63
_PG_IDENT_MAX = 63
def _safe_schema(site_name: str) -> str:
    if not site_name:
        return "s_default"
    s = (site_name or "").strip().lower()
    s = s.replace("-", "_")
    s = re.sub(r"[^a-z0-9_]", "_", s)
    s = re.sub(r"_+", "_", s)
    if not s:
        s = "s_default"
    if s[0].isdigit():
        s = "s_" + s
    s = s[:_PG_IDENT_MAX]
    return s

@app.route("/create_form/site/<site>", methods=["GET"])
@admin_required
def create_form_site(site):
    conn = get_conn(); c = conn.cursor()
    try:
        c.execute(
            """
            SELECT name, site_name, schema_json, COALESCE(description,''::TEXT)
              FROM form_defs
             WHERE site_name=%s
             LIMIT 1
            """,
            (site,),
        )
        row = c.fetchone()
        if not row:
            return redirect(url_for("create_form_new", enter=1))

        form_name, site_name, schema_json, form_desc = row

        submissions = []
        try:
            schema_name = _safe_schema(site_name)
            c.execute(f"SET search_path TO {schema_name}")
            c.execute(
                "SELECT id, user_id, data, status, created_at FROM submissions ORDER BY id DESC LIMIT 50"
            )
            submissions = c.fetchall()
        except Exception:
            pass
    finally:
        conn.close()

    return render_template(
        "create_form.html",
        form_name=form_name,
        form_desc=form_desc,
        site_name=site_name,
        schema_json=schema_json,
        submissions=submissions,
    )

# ========= 创建页按钮微调注入 =========
@app.after_request
def _tweak_create_form_image_button(resp):
    try:
        if request.path.startswith("/create_form") and resp.content_type.startswith("text/html"):
            html = resp.get_data(as_text=True)
            inject_css = """
<style id="cehs-imgbtn-style">
  .cehs-imgbtn{
    display:inline-flex !important;
    align-items:center !important;
    gap:6px !important;
    padding:8px 12px !important;
    border-radius:10px !important;
    background:#fff !important;
    color:#111 !important;
    border:1px solid #e5e7eb !important;
    box-shadow:none !important;
    line-height:1 !important;
    height:auto !important;
    font-weight:600 !important;
  }
  .cehs-imgbtn:hover{ background:#fafafa !important; border-color:#d1d5db !important; }
  .cehs-imgbtn svg, .cehs-imgbtn img{ width:18px !important; height:18px !important; flex:0 0 auto !important; }
</style>
"""
            inject_js = """
<script>
(function(){
  function normalizeImageButtons(){
    var btns = Array.from(document.querySelectorAll('button, .btn, .tool, [role="button"]'));
    btns.forEach(function(el){
      var txt = (el.textContent || '').trim();
      if (txt === '图片' && !el.classList.contains('cehs-imgbtn')) {
        el.classList.add('cehs-imgbtn');
        try {
          el.style.display = 'inline-flex';
          el.style.flexDirection = 'row';
          el.style.alignItems = 'center';
          el.style.gap = '6px';
        } catch(_) {}
        var hasTextNode = Array.from(el.childNodes).some(function(n){
          return (n.nodeType === 3 && n.textContent.trim()) || (n.nodeType === 1 && (n.textContent||'').trim());
        });
        if (!hasTextNode) {
          var s = document.createElement('span');
          s.textContent = '图片';
          el.appendChild(s);
        }
      }
    });
  }
  if (document.readyState === 'loading') {
    document.addEventListener('DOMContentLoaded', normalizeImageButtons);
  } else {
    normalizeImageButtons();
  }
  var mo = new MutationObserver(normalizeImageButtons);
  mo.observe(document.body, {childList:true, subtree:true});
})();
</script>
"""
            if "</body>" in html:
                html = html.replace("</body>", inject_css + inject_js + "</body>")
                resp.set_data(html)
    except Exception:
        pass
    return resp

# ========= 管理端上传/删除/主题保存 =========
@app.route("/site/<site_name>/admin")
@admin_required
def site_admin(site_name):
    conn = get_conn(); c = conn.cursor()
    c.execute("SELECT schema_json FROM form_defs WHERE site_name=%s", (site_name,))
    row = c.fetchone(); conn.close()
    schema = row[0] if row and isinstance(row[0], dict) else (json.loads(row[0]) if row and row[0] else {})
    brand_light, brand_dark, theme_mode = _read_theme(schema)

    return render_template(
        "dynamic_admin.html",
        site_name=site_name,
        brand_light=brand_light,
        brand_dark=brand_dark,
        theme_mode=theme_mode,
    )

@app.route("/site/<site_name>/admin/api/upload_asset", methods=["POST"])
@admin_required
def api_upload_asset(site_name):
    f = request.files.get("file")
    if not f or not f.filename:
        return jsonify({"ok": False, "error": "未选择文件"}), 400

    # 读取 schema 的白名单
    conn = get_conn(); c = conn.cursor()
    c.execute("SELECT schema_json FROM form_defs WHERE site_name=%s", (site_name,))
    row = c.fetchone(); conn.close()
    schema = row[0] if row and isinstance(row[0], dict) else (json.loads(row[0]) if row and row[0] else {})
    upload_cfg = (schema.get("upload") or (schema.get("settings") or {}).get("upload") or {})
    allowed = set(x.strip().lower() for x in str(upload_cfg.get("allowed_file_types","")).split(",") if x.strip())

    ext = Path(f.filename).suffix.lower().lstrip(".")
    if allowed and ext not in allowed:
        return jsonify({"ok": False, "error": f"不允许的文件类型: .{ext}"}), 400

    filename = secure_filename(f.filename)
    folder = os.path.join(app.config["UPLOAD_FOLDER"], site_name)
    os.makedirs(folder, exist_ok=True)
    path = os.path.join(folder, filename)
    f.save(path)
    url = url_for("site_uploaded_file", site_name=site_name, filename=filename, _external=True)
    return jsonify({"ok": True, "filename": filename, "url": url})

@app.route("/site/<site_name>/admin/api/delete_asset", methods=["POST"])
@admin_required
def api_delete_asset(site_name):
    name = (request.json or {}).get("filename", "")
    safe = secure_filename(name)
    if not safe or "/" in name or safe != name:
        return jsonify({"ok": False, "error": "非法文件名"}), 400

    # ✅ 使用站点目录
    folder = os.path.join(app.config["UPLOAD_FOLDER"], site_name)
    path = os.path.join(folder, safe)

    if os.path.exists(path):
        os.remove(path)
    return jsonify({"ok": True})

@app.route("/site/<site_name>/admin/api/save_theme_bg", methods=["POST"])
@admin_required
def api_save_theme_bg(site_name):
    payload = request.get_json() or {}
    t = payload.get("theme") or {}

    brand_light = (t.get("brand_light") or t.get("brand") or "").strip()
    brand_dark  = (t.get("brand_dark")  or "").strip()
    mode        = (t.get("mode") or t.get("theme_mode") or "auto").lower()
    if mode not in ("light","dark","auto"):
        mode = "auto"

    conn = get_conn(); c = conn.cursor()
    try:
        # 1) 读库得到 schema
        c.execute("SELECT schema_json FROM form_defs WHERE site_name=%s", (site_name,))
        row = c.fetchone()
        schema = row[0] if row and isinstance(row[0], dict) else (json.loads(row[0]) if row and row[0] else {})

        # 2) 写 header.title_image（以及可选位置）
        bg = payload.get("bg")
        bg_pos = payload.get("bg_position")
        header = schema.get("header") or {}
        if bg:
            header["title_image"] = bg
        if bg_pos:
            header["title_image_pos"] = bg_pos
        schema["header"] = header

        # 3) 仍然按你现有逻辑写 theme（brand_light/brand_dark/mode）
        theme = schema.get("theme") or {}
        theme["brand"] = brand_light or theme.get("brand", "")
        theme["brand_dark"] = brand_dark or theme.get("brand_dark", "")
        theme["mode"] = mode
        schema["theme"] = theme

        # 4) 回写
        c.execute("UPDATE form_defs SET schema_json=%s WHERE site_name=%s", (Json(schema), site_name))
        conn.commit()

        return jsonify({"ok": True})
    except Exception as e:
        conn.rollback()
        return jsonify({"ok": False, "error": str(e)}), 500
    finally:
        conn.close()

@app.route("/site/<site_name>/admin/api/delete", methods=["POST"])
@admin_required
def api_delete_submission(site_name):
    payload = request.get_json() or {}
    sub_id = int(payload.get("id") or 0)
    if sub_id <= 0:
        return jsonify({"ok": False, "error": "参数错误"}), 400

    schema = _safe_schema(site_name)
    conn = get_conn(); c = conn.cursor()
    try:
        c.execute(f'SET search_path TO "{schema}", public')
        c.execute("DELETE FROM submissions WHERE id=%s", (sub_id,))
        conn.commit()
        return jsonify({"ok": True})
    except Exception as e:
        conn.rollback()
        return jsonify({"ok": False, "error": str(e)}), 500
    finally:
        conn.close()

def _extract_columns_from_schema(schema: dict):
    cols = []
    if not isinstance(schema, dict):
        return cols

    # ---------- helpers ----------
    def _to_text(sval) -> str:
        if not isinstance(sval, str):
            return ""
        sval = re.sub(r"<[^>]+>", "", sval)
        return sval.strip()

    def _has_cjk(text: str) -> bool:
        """是否包含中文（CJK）"""
        return bool(re.search(r"[\u4e00-\u9fff]", str(text or "")))

    def pick_label(f: dict) -> str:
        # 覆盖常见构建器字段名
        for cand in (
            f.get("label"), f.get("title"), f.get("text"), f.get("name"),
            f.get("placeholder"), f.get("question"), f.get("displayName"),
            f.get("desc"), f.get("description"),
            (f.get("ui") or {}).get("label"), (f.get("ui") or {}).get("title"),
            (f.get("props") or {}).get("label"), (f.get("props") or {}).get("title"),
            (f.get("meta") or {}).get("label"), (f.get("meta") or {}).get("title"),
        ):
            t = _to_text(cand) if cand else ""
            if t:
                return t
        # i18n / 富文本兜底
        for key in ("i18n", "labelHTML", "label", "title", "question"):
            obj = f.get(key)
            if isinstance(obj, dict):
                for lang_key in ("zh-CN","zh_CN","zh-cn","zh","text","title","label","question","en"):
                    t = _to_text(obj.get(lang_key) or "")
                    if t:
                        return t
        return ""

    # 深度遍历 schema，尽可能找出字段节点
    CAND_KEYS = {"fields","questions","items","components","children",
                 "body","rows","columns","pages","formItems","list",
                 "properties","elements","schema"}
    seen_keys = set()

    def iter_fields(node):
        if isinstance(node, dict):
            # 当前节点本身可能就是字段
            if any(k in node for k in ("key","id","name")) and any(k in node for k in ("label","title","text","question","displayName")):
                yield node
            for k, v in node.items():
                if k in CAND_KEYS or isinstance(v, (list, dict)):
                    yield from iter_fields(v)
        elif isinstance(node, list):
            for it in node:
                yield from iter_fields(it)

    for f in iter_fields(schema):
        if not isinstance(f, dict):
            continue
        key = f.get("key") or f.get("id") or f.get("name")
        if not key or key in seen_keys:
            continue
        label = pick_label(f)
        # 只保留“含中文标题”的列（你现在的需求）
        if not label or not _has_cjk(label):
            continue
        type_ = f.get("type") or (f.get("ui") or {}).get("type") or ""
        cols.append({"key": str(key), "label": label, "type": str(type_)})
        seen_keys.add(key)

    return cols

@app.route("/site/<site_name>/admin/api/responses")
@admin_required
def api_responses(site_name):
    # 这段原始代码里有未定义变量（c/results），保留原样不改动，避免影响你其它逻辑的期待
    # 建议你使用 /site/<site_name>/admin/api/submissions 这个端点
    return _api_list_responses(site_name)

@app.route("/site/<site_name>/admin/api/submissions")  # 兼容旧别名
@admin_required
def api_responses_alias(site_name):
    return _api_list_responses(site_name)

def _api_list_responses(site_name: str):
    q = (request.args.get("q") or "").strip()
    schema = _safe_schema(site_name)

    # 读提交数据
    conn = get_conn(); c = conn.cursor()
    try:
        c.execute(f'SET search_path TO "{schema}", public')
        if q:
            c.execute("""
                SELECT id, data, status, review_comment, created_at
                  FROM submissions
                 WHERE data::text ILIKE %s
                 ORDER BY id DESC
                 LIMIT 500
            """, (f"%{q}%",))
        else:
            c.execute("""
                SELECT id, data, status, review_comment, created_at
                  FROM submissions
                 ORDER BY id DESC
                 LIMIT 500
            """)
        rows = c.fetchall()
    except Exception as e:
        conn.rollback()
        return jsonify({"ok": False, "error": str(e)}), 500
    finally:
        conn.close()

    # 数据项：统一做 UTF-8 归一化
    items = []
    for rid, d, status, review, created in rows:
        data = d if isinstance(d, dict) else (json.loads(d) if d else {})
        data = _normalize_obj(data)
        items.append({
            "id": rid,
            "status": _maybe_fix_encoding(status or "待审核"),
            "review_comment": _maybe_fix_encoding(review or ""),
            "created_at": str(created) if created else "",
            "data": data,
        })

    # 读 schema，生成“中文列头”
    conn2 = get_conn(); c2 = conn2.cursor()
    try:
        c2.execute(f'SET search_path TO "{schema}", public')
        c2.execute("SELECT schema_json FROM form_defs WHERE site_name=%s", (site_name,))
        row = c2.fetchone()
    except Exception:
        row = None
    finally:
        conn2.close()

    schema_json = row[0] if (row and isinstance(row[0], dict)) else (json.loads(row[0]) if row and row[0] else {})
    columns = _extract_columns_from_schema(schema_json) if schema_json else []

    cleaned = []
    for c in (columns or []):
        key = c.get("key") or c.get("name") or c.get("id")
        if not key: continue
        label = c.get("label") or c.get("title") or c.get("text") or key
        cleaned.append({
            "key": str(key),
            "label": _maybe_fix_encoding(str(label)),
            "type": c.get("type", "")
        })
    title_map = {c["key"]: c["label"] for c in cleaned if c.get("key")}

    return jsonify({"ok": True, "items": items, "columns": cleaned, "titleMap": title_map})

# ========= 公共页回退模板 =========
PUBLIC_FORM_HTML = """
<!doctype html>
<meta charset="utf-8">
<title>{{ form_name or site_name }} - 表单填写</title>
<style>
  :root{ --accent: {{ brand|default('#2563eb') }}; }
  body{
    font-family:Microsoft YaHei,Arial;
    padding:20px;
    background:#f7f7f7;
  }
  .box{
    max-width:720px;margin:0 auto;background:#fff;padding:20px;border-radius:12px;
    box-shadow:0 2px 12px rgba(0,0,0,.08)
  }
  h1{margin:0 0 8px 0}
  .desc{color:#666;margin:2px 0 10px 0;line-height:1.6}
  .row{margin:12px 0}
  input,select,textarea{width:100%;padding:10px;border:1px solid #ddd;border-radius:8px}
  button{padding:12px 18px;border:0;border-radius:8px;background:var(--accent);color:#fff;font-weight:700;cursor:pointer}
  .ok{background:var(--accent)}
</style>
<div class="box">
  <h1>{{ form_name or site_name }}</h1>
  {% if form_desc %}<div class="desc">{{ form_desc|safe }}</div>{% endif %}
  <form method="post" enctype="multipart/form-data">
    {% for f in fields %} 
      {% set input_name = f.key or f.id %}
      <div class="row">
        <label><strong>{{ f.label }}</strong></label><br>
        {% if f.type in ['text','email','number','date','time'] %}
          <input name="{{ input_name }}" type="{{ f.type }}" {{ 'required' if f.required else '' }}>
        {% elif f.type=='textarea' %}
          <textarea name="{{ input_name }}" rows="3" {{ 'required' if f.required else '' }}></textarea>
        {% elif f.type in ['radio','checkbox'] %}
          {% for opt in f.options or [] %}
            <label><input type="{{ 'radio' if f.type=='radio' else 'checkbox' }}" name="{{ input_name }}{% if f.type=='checkbox' %}[]{% endif %}" value="{{ opt }}"> {{ opt }}</label>&nbsp;&nbsp;
          {% endfor %}
        {% elif f.type=='select' %}
          <select name="{{ input_name }}">
            {% for opt in f.options or [] %}<option value="{{ opt }}">{{ opt }}</option>{% endfor %}
          </select>
        {% elif f.type=='file' %}
          <input type="file" name="{{ input_name }}">
        {% else %}
          <input name="{{ input_name }}" type="text" {{ 'required' if f.required else '' }}>
        {% endif %}
      </div>
    {% endfor %}
    <button type="submit" class="ok">提交</button>
  </form>
</div>
"""

# ========= 公开提交成功 =========
PUBLIC_SUCCESS_HTML = """
<!doctype html>
<meta charset="utf-8">
<title>提交成功</title>
<style>
  body{font-family:Microsoft YaHei,Arial;background:#f6f7fb;padding:30px}
  .box{max-width:680px;margin:8vh auto;background:#fff;border:1px solid #e5e7eb;border-radius:14px;box-shadow:0 10px 30px rgba(0,0,0,.08);padding:22px}
  .btn{display:inline-block;margin-right:10px;padding:10px 14px;border:0;border-radius:10px;background:#2563eb;color:#fff;font-weight:800;text-decoration:none}
  .btn.ghost{background:#fff;color:#111;border:1px solid #e5e7eb}
  .pill{display:inline-block;padding:4px 10px;border-radius:999px;font-weight:700;font-size:12px}
  .pill.good{background:#dcfce7;color:#15803d}
  .pill.bad{background:#fee2e2;color:#b91c1c}
  .pill.wait{background:#e5e7eb;color:#111827}
  .modal{position:fixed;inset:0;display:none;align-items:center;justify-content:center;background:rgba(0,0,0,.45);z-index:50;padding:16px}
  .modal.show{display:flex}
  .modal .box2{background:#fff;width:min(520px,100%);border-radius:14px;box-shadow:0 10px 30px rgba(0,0,0,.12);padding:16px;border:1px solid #e5e7eb}
  input{width:100%;padding:10px;border:1px solid #ddd;border-radius:10px}
</style>
<div class="box">
  <h2>提交成功 🎉</h2>
  <p>我们已收到你的提交，请稍后到“查看状态”里查看审核结果。</p>
  <p>
    <a href="{{ public_url }}" class="btn ghost">返回表单主页</a>
    <a href="javascript:void(0)" id="btnCheck" class="btn">查看状态</a>
  </p>
</div>

<div class="modal" id="statusModal" aria-hidden="true">
  <div class="box2">
    <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:10px">
      <strong>查看审核状态</strong>
      <a href="javascript:void(0)" id="closeStatus" class="btn ghost">关闭</a>
    </div>
    <div style="display:flex;gap:8px;align-items:center;margin-bottom:10px;flex-wrap:wrap">
      <input id="statusName" placeholder="请输入姓名">
      <a href="javascript:void(0)" id="goQuery" class="btn">查看</a>
    </div>
    <div id="statusResult" style="font-size:14px;color:#111"></div>
  </div>
</div>

<script>
  const btnCheck = document.getElementById('btnCheck');
  const modal = document.getElementById('statusModal');
  const closeStatus = document.getElementById('closeStatus');
  const goQuery = document.getElementById('goQuery');
  const statusName = document.getElementById('statusName');
  const statusResult = document.getElementById('statusResult');

  btnCheck?.addEventListener('click', ()=>{ modal.classList.add('show'); statusName.focus(); });
  closeStatus?.addEventListener('click', ()=> modal.classList.remove('show'));
  modal?.addEventListener('click', (e)=>{ if(e.target===modal) modal.classList.remove('show'); });

  async function doQuery(){
    const name = (statusName.value||'').trim();
    if(!name){ statusResult.innerHTML='<span style="color:#b91c1c">请输入姓名</span>'; return; }
    statusResult.textContent='查询中…';
    try{
      const res = await fetch('/site/{{ site_name }}/status_query?name='+encodeURIComponent(name));
      const j = await res.json();
      if(!res.ok || !j.ok){ statusResult.innerHTML='查询失败：'+(j.error||res.status); return; }
      if(!j.found){ statusResult.innerHTML='没有找到相关记录'; return; }
      let pill = '<span class="pill wait">待审核</span>';
      if(j.status==='已通过') pill='<span class="pill good">已通过</span>';
      else if(j.status==='未通过') pill='<span class="pill bad">未通过</span>';
      const cmt = j.review_comment ? ('<div style="margin-top:6px;color:#374151">说明：'+j.review_comment+'</div>') : '';
      const when = j.created_at ? ('<div style="margin-top:6px;color:#6b7280">提交时间：'+j.created_at+'</div>') : '';
      statusResult.innerHTML = '<div>最新状态：'+pill+'</div>'+cmt+when;
    }catch(e){
      statusResult.innerHTML='查询失败';
    }
  }
  goQuery?.addEventListener('click', doQuery);
  statusName?.addEventListener('keydown', e=>{ if(e.key==='Enter'){ e.preventDefault(); doQuery(); } });
</script>
"""

@app.route("/site/<site_name>/admin/api/review", methods=["POST"])
@admin_required
def api_review(site_name):
    payload = request.get_json() or {}
    sub_id = int(payload.get("id") or 0)
    status = (payload.get("status") or "").strip()
    review_comment = (payload.get("review_comment") or "").strip()

    if sub_id <= 0 or status not in ("已通过", "未通过", "待审核"):
        return jsonify({"ok": False, "error": "参数错误"}), 400

    schema = _safe_schema(site_name)
    conn = get_conn(); c = conn.cursor()
    try:
        c.execute(f'SET search_path TO "{schema}", public')
        c.execute("UPDATE submissions SET status=%s, review_comment=%s WHERE id=%s",
                  (status, review_comment, sub_id))
        conn.commit()
        return jsonify({"ok": True})
    except Exception as e:
        conn.rollback()
        return jsonify({"ok": False, "error": str(e)}), 500
    finally:
        conn.close()

@app.route("/site/<site_name>/admin/api/send_mail", methods=["POST"])
@admin_required
def api_send_mail(site_name):
    payload = request.get_json() or {}
    sub_id = int(payload.get("id") or 0)
    schema = _safe_schema(site_name)

    conn = get_conn(); c = conn.cursor()
    c.execute(f'SET search_path TO "{schema}", public')
    c.execute("SELECT data, status, review_comment FROM submissions WHERE id=%s", (sub_id,))
    row = c.fetchone()
    conn.close()
    if not row:
        return jsonify({"ok": False, "error": "记录不存在"}), 404

    data = row[0] if isinstance(row[0], dict) else (json.loads(row[0]) if row[0] else {})
    status = row[1] or "待审核"
    review_comment = row[2] or ""
    to_email = _extract_email(data)

    if not to_email:
        return jsonify({"ok": False, "error": "记录中未找到邮箱字段"}), 400
    if not (SENDER_EMAIL and SENDER_PASSWORD and SMTP_SERVER):
        return jsonify({"ok": False, "error": "SMTP 未配置"}), 400

    subject = f"[{site_name}] 审核结果通知：{status}"
    body = f"您好！您的申请已审核：{status}\n审核说明：{review_comment or '（无）'}\n\n— {site_name}"
    msg = MIMEText(body, "plain", "utf-8")
    msg["Subject"] = Header(subject, "utf-8")
    msg["From"] = formataddr((str(Header(site_name, "utf-8")), SENDER_EMAIL))
    msg["To"] = to_email

    try:
        smtp = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
        smtp.starttls()
        smtp.login(SENDER_EMAIL, SENDER_PASSWORD)
        smtp.sendmail(SENDER_EMAIL, [to_email], msg.as_string())
        smtp.quit()
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

# ========= 公开表单 GET =========
@app.route("/f/<site_name>", methods=["GET"])
def public_form(site_name):
    conn = get_conn(); c = conn.cursor()
    c.execute("SELECT name, schema_json, COALESCE(description,'') FROM form_defs WHERE site_name=%s", (site_name,))
    row = c.fetchone()
    conn.close()
    if not row:
        abort(404)

    form_title = row[0]
    schema = row[1] if isinstance(row[1], dict) else json.loads(row[1] or "{}")
    db_desc = (row[2] or "").strip()
    form_desc_html = (
        (schema.get("descHTML") or schema.get("desc") or schema.get("description"))
        or db_desc
    )

    # ✅ 从 schema 里取主题
    theme = schema.get('theme') or {}

    # 主题/外观
    brand_light = (theme.get("brand_light") or theme.get("brand") or "#2563eb").strip()
    brand_dark  = (theme.get("brand_dark")  or theme.get("brand") or "#0ea5e9").strip()
    theme_mode  = (theme.get("mode") or theme.get("theme_mode") or theme.get("appearance") or "auto").lower()
    if theme_mode not in ("light", "dark", "auto"):
        theme_mode = "auto"

    # 上传配置
    upload_cfg = (schema.get("upload") or (schema.get("settings") or {}).get("upload") or {})
    upload_max_files = int(upload_cfg.get("max_files") or 3)
    allowed = set(
        x.strip().lower()
        for x in str(upload_cfg.get("allowed_file_types", "")).split(",")
        if x.strip()
    )

    # 字段（仅去掉描述类键，其他不动）
    raw_fields = schema.get("fields") or []
    if not isinstance(raw_fields, list):
        raw_fields = []
    def _strip_desc(f):
        g = dict(f or {})
        for k in ("desc","descHTML","description","help","helpText"):
            g.pop(k, None)
        return g
    clean_fields = [_strip_desc(f) for f in raw_fields]

    # 兜底字段（仅用于 TemplateNotFound 回退）
    fields_fallback = []
    for f in clean_fields:
        fields_fallback.append({
            "label": f.get("labelHTML") or f.get("label") or f.get("key",""),
            "type": (f.get("type") or "text"),
            "key":  (f.get("key") or f.get("id")),
            "options": f.get("options") or [],
            "required": bool(f.get("required", False)),
        })

    try:
        return render_template(
            "public_form.html",
            site_name=site_name,
            form_title=form_title,
            form_desc=form_desc_html,
            fields=clean_fields,
            brand_light=brand_light,
            brand_dark=brand_dark,
            theme_mode=theme_mode,
            has_file=any((f.get("type") or "").lower() == "file" for f in clean_fields),
            upload_max_files=upload_max_files,
            schema_json=json.dumps(schema, ensure_ascii=False),  # ← 新增
        )
    except TemplateNotFound:
        # 简易回退模板
        brand = brand_dark if theme_mode == "dark" else brand_light
        return render_template_string(
            PUBLIC_FORM_HTML,
            site_name=site_name,
            form_name=form_title,
            form_desc=form_desc_html,
            fields=fields_fallback,
            brand=brand,
        )

def _read_theme(schema: dict):
    theme = schema.get("theme") if isinstance(schema.get("theme"), dict) else {}
    brand_single = (theme.get("brand") or "").strip()
    brand_light = (theme.get("brand_light") or brand_single or "#2563eb").strip()
    brand_dark  = (theme.get("brand_dark")  or brand_single or brand_light or "#0ea5e9").strip()
    mode = (theme.get("mode") or theme.get("theme_mode") or theme.get("appearance") or "auto").lower()
    if mode not in ("light", "dark", "auto"):
        mode = "auto"
    return brand_light, brand_dark, mode

# ========= 公开表单 POST =========
@app.route("/f/<site_name>", methods=["POST"])
def public_submit(site_name):
    schema_name = _safe_schema(site_name)
    conn = get_conn(); c = conn.cursor()

    # 读 schema（上传上限）
    c.execute("SELECT schema_json FROM form_defs WHERE site_name=%s", (site_name,))
    row = c.fetchone()
    if not row:
        conn.close()
        return "not found", 404
    schema = row[0] if isinstance(row[0], dict) else json.loads(row[0] or "{}")

    upload_cfg = (schema.get("upload") or (schema.get("settings") or {}).get("upload") or {})
    max_files = int(upload_cfg.get("max_files") or 3)
    allowed = set(
        x.strip().lower()
        for x in str(upload_cfg.get("allowed_file_types", "")).split(",")
        if x.strip()
    )

    # 非文件字段
    data = request.form.to_dict()

    # 站点专属目录
    site_folder = os.path.join(app.config["UPLOAD_FOLDER"], site_name)
    os.makedirs(site_folder, exist_ok=True)

    for field_key in request.files:
        files = request.files.getlist(field_key)
        saved_urls = []
        for f in files[:max_files]:
            if not f or not f.filename:
                continue
            ext = Path(f.filename).suffix.lower().lstrip(".")
            if allowed and ext not in allowed:
                continue
            uniq = f"{int(time())}_{uuid4().hex[:8]}_{secure_filename(f.filename)}"
            abs_path = os.path.join(site_folder, uniq)
            f.save(abs_path)
            saved_urls.append(f"/site/{site_name}/uploads/{uniq}")
        if saved_urls:
            data[field_key] = saved_urls

    # 写入“namespaced”表
    try:
        c.execute(f'CREATE SCHEMA IF NOT EXISTS "{schema_name}"')
        c.execute(f"""
            CREATE TABLE IF NOT EXISTS "{schema_name}".submissions (
                id SERIAL PRIMARY KEY,
                user_id INT,
                data JSONB,
                status TEXT DEFAULT '待审核',
                review_comment TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        c.execute(f'INSERT INTO "{schema_name}".submissions (data, status) VALUES (%s, %s) RETURNING id',
                  (json.dumps(data, ensure_ascii=False), '待审核'))
        new_id = c.fetchone()[0]
        conn.commit()
    except Exception as e:
        conn.rollback(); conn.close()
        return f"提交失败：{e}", 500
    finally:
        conn.close()

    # 成功页（公共）
    return render_template_string(
        PUBLIC_SUCCESS_HTML,
        site_name=site_name,
        home_url=url_for("index"),
        public_url=url_for("public_form", site_name=site_name)
    )

# ========= 站点内上传文件访问 =========
@app.route("/site/<site_name>/uploads/<path:filename>")
def site_uploaded_file(site_name, filename):
    folder = os.path.join(app.config.get("UPLOAD_FOLDER", "uploads"), site_name)
    return send_from_directory(folder, filename, as_attachment=False)

# ========= 状态查询 =========
@app.route("/site/<site_name>/status_query")
def public_status_query(site_name):
    name = (request.args.get("name") or "").strip()
    if not name:
        return jsonify({"ok": False, "error": "缺少姓名"}), 400

    schema = _safe_schema(site_name)
    conn = get_conn(); c = conn.cursor()
    try:
        c.execute(f'SET search_path TO "{schema}", public')
        c.execute("""
            SELECT id, data, status, review_comment, created_at
            FROM submissions
            WHERE data::text ILIKE %s
            ORDER BY created_at DESC
            LIMIT 1
        """, (f"%{name}%",))
        row = c.fetchone()
        if not row:
            return jsonify({"ok": True, "found": False})
        data = row[1] if isinstance(row[1], dict) else (json.loads(row[1]) if row[1] else {})
        return jsonify({
            "ok": True, "found": True,
            "id": row[0],
            "status": row[2] or "待审核",
            "review_comment": row[3] or "",
            "created_at": str(row[4]) if row[4] else "",
            "data": data
        })
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500
    finally:
        conn.close()

# 兼容旧 URL（不分站点）
@app.route('/uploads/<path:filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

# ========= 删除与导出 =========
@app.route("/form/<int:form_id>/delete/<int:sub_id>", methods=["GET","POST"])
@admin_required
def delete_submission(form_id, sub_id):
    conn = get_conn(); c = conn.cursor()
    c.execute("SELECT id, site_name, db_url, created_by FROM form_defs WHERE id=%s", (form_id,))
    row = c.fetchone()
    if not row:
        conn.close()
        return jsonify({"success":False,"message":"表单不存在"})
    site_name = row[1]
    schema_name = row[2]
    c.execute(f"SET search_path TO {schema_name}")
    c.execute("DELETE FROM submissions WHERE id=%s", (sub_id,))
    conn.commit(); conn.close()
    if request.method=="POST" and request.is_json:
        return jsonify({"success":True})
    return redirect(url_for("site_admin", site_name=site_name))

@app.route("/form/<int:form_id>/delete", methods=["POST"])
@admin_required
def delete_form(form_id):
    user_id = session.get("user_id")
    conn = get_conn(); c = conn.cursor()
    c.execute("SELECT id, site_name, db_url, created_by FROM form_defs WHERE id=%s", (form_id,))
    row = c.fetchone()
    if not row:
        conn.close()
        return jsonify({"ok": False, "error": "表单不存在"}), 404
    if row[3] != user_id:
        conn.close()
        return jsonify({"ok": False, "error": "无权限"}), 403

    schema_name = row[2]
    try:
        c.execute(f"DROP SCHEMA IF EXISTS {schema_name} CASCADE")
        c.execute("DELETE FROM form_defs WHERE id=%s", (form_id,))
        conn.commit()
        ok = True; err = None
    except Exception as e:
        conn.rollback()
        ok = False; err = str(e)
    finally:
        conn.close()

    if not ok:
        return jsonify({"ok": False, "error": err}), 500
    return jsonify({"ok": True})

@app.route("/site/<site_name>/preview", methods=["GET", "POST"])
@admin_required
def preview_form(site_name):
    """
    预览当前未发布的表单：渲染 public_form.html，但不落库、不允许提交
    - POST: 优先使用请求里带来的 schema_json/form_name/form_desc
    - GET: 退化为预览数据库里已保存的表单（等价 /f/<site_name>）
    """
    # 读取 schema / 标题 / 描述
    schema = {}
    form_title = site_name
    form_desc_html = ""

    if request.method == "POST":
        raw = request.form.get("schema_json")
        if not raw and request.is_json:
            raw = (request.get_json(silent=True) or {}).get("schema_json")
        try:
            schema = json.loads(raw) if isinstance(raw, str) else (raw or {})
        except Exception:
            schema = {}

        form_title = (request.form.get("form_name") or
                      schema.get("title") or schema.get("name") or site_name)
        form_desc_html = (request.form.get("form_desc") or
                          schema.get("descHTML") or schema.get("desc") or schema.get("description") or "")
    else:
        # 和 /f/<site_name> 一致：从数据库读
        conn = get_conn(); c = conn.cursor()
        c.execute("SELECT name, schema_json, COALESCE(description,'') FROM form_defs WHERE site_name=%s", (site_name,))
        row = c.fetchone(); conn.close()
        if not row:
            abort(404)
        form_title = row[0]
        schema = row[1] if isinstance(row[1], dict) else json.loads(row[1] or "{}")
        db_desc = (row[2] or "").strip()
        form_desc_html = (schema.get("descHTML") or schema.get("desc") or schema.get("description") or db_desc)

    # 主题/外观（与 public_form 保持一致）
    theme = schema.get("theme") or {}
    brand_light = (theme.get("brand_light") or theme.get("brand") or "#2563eb").strip()
    brand_dark  = (theme.get("brand_dark")  or theme.get("brand") or "#0ea5e9").strip()
    theme_mode  = (theme.get("mode") or theme.get("theme_mode") or theme.get("appearance") or "auto").lower()
    if theme_mode not in ("light","dark","auto"):
        theme_mode = "auto"

    upload_cfg = (schema.get("upload") or (schema.get("settings") or {}).get("upload") or {})
    upload_max_files = int(upload_cfg.get("max_files") or 3)

    # 字段清洗（同 public_form）
    raw_fields = schema.get("fields") or []
    if not isinstance(raw_fields, list):
        raw_fields = []
    def _strip_desc(f):
        g = dict(f or {})
        for k in ("desc","descHTML","description","help","helpText"):
            g.pop(k, None)
        return g
    clean_fields = [_strip_desc(f) for f in raw_fields]

    # 模板渲染
    try:
        return render_template(
            "public_form.html",
            site_name=site_name,
            form_title=form_title,
            form_desc=form_desc_html,
            fields=clean_fields,
            brand_light=brand_light,
            brand_dark=brand_dark,
            theme_mode=theme_mode,
            has_file=any((f.get("type") or "").lower() == "file" for f in clean_fields),
            upload_max_files=upload_max_files,
            # 传个标记给前端，如有用可用它做定制
            preview_mode=True,
            schema_json=json.dumps(schema, ensure_ascii=False),  # ← 新增
        )
    except TemplateNotFound:
        brand = brand_dark if theme_mode == "dark" else brand_light
        return render_template_string(
            PUBLIC_FORM_HTML,
            site_name=site_name,
            form_name=form_title,
            form_desc=form_desc_html,
            fields=[{
                "label": f.get("labelHTML") or f.get("label") or f.get("key",""),
                "type": (f.get("type") or "text"),
                "key":  (f.get("key") or f.get("id")),
                "options": f.get("options") or [],
                "required": bool(f.get("required", False)),
            } for f in clean_fields],
            brand=brand,
        )

# === 预览：不需要 site_name，直接按传入 schema 渲染公开页 ===
@app.post("/preview")
@admin_required
def preview_inline():
    form_name   = (request.form.get("form_name") or "").strip()
    form_desc   = (request.form.get("form_desc") or "").strip()
    schema_json = request.form.get("schema_json") or "{}"

    try:
        schema = json.loads(schema_json) if isinstance(schema_json, str) else (schema_json or {})
    except Exception:
        schema = {}

    # 主题参数
    theme = schema.get("theme") or {}
    brand_light = (theme.get("brand_light") or theme.get("brand") or "#2563eb").strip()
    brand_dark  = (theme.get("brand_dark")  or theme.get("brand") or "#0ea5e9").strip()
    theme_mode  = (theme.get("mode") or theme.get("theme_mode") or theme.get("appearance") or "auto").lower()
    if theme_mode not in ("light","dark","auto"):
        theme_mode = "auto"

    # 上传配置（仅用于渲染控件外观；预览不真正上传/提交）
    upload_cfg = (schema.get("upload") or (schema.get("settings") or {}).get("upload") or {})
    upload_max_files = int(upload_cfg.get("max_files") or 3)

    # 字段清洗（去掉描述类键）
    raw_fields = schema.get("fields") or []
    if not isinstance(raw_fields, list):
        raw_fields = []
    def _strip_desc(f):
        g = dict(f or {})
        for k in ("desc","descHTML","description","help","helpText"):
            g.pop(k, None)
        return g
    clean_fields = [_strip_desc(f) for f in raw_fields]

    # 回退字段（用于 fallback 模板）
    fields_fallback = []
    for f in clean_fields:
        fields_fallback.append({
            "label": f.get("labelHTML") or f.get("label") or f.get("key",""),
            "type":  (f.get("type") or "text"),
            "key":   (f.get("key") or f.get("id")),
            "options": f.get("options") or [],
            "required": bool(f.get("required", False)),
        })

    # 用和公开页相同的模板渲染；site_name 给个占位，防止误写库
    try:
        return render_template(
            "public_form.html",
            site_name="__preview__",                       # 占位，表单提交会 404，不会写库
            form_title=form_name or schema.get("name") or "预览",
            form_desc=form_desc or (schema.get("descHTML") or schema.get("desc") or schema.get("description")),
            fields=clean_fields,
            brand_light=brand_light,
            brand_dark=brand_dark,
            theme_mode=theme_mode,
            has_file=any((f.get("type") or "").lower() == "file" for f in clean_fields),
            upload_max_files=upload_max_files,
            schema_json=json.dumps(schema, ensure_ascii=False),  # ← 新增

            preview_mode=True                              # 模板可选识别（如需禁用提交）
        )
    except TemplateNotFound:
        # 使用你已有的回退 HTML
        brand = brand_dark if theme_mode == "dark" else brand_light
        return render_template_string(
            PUBLIC_FORM_HTML,
            site_name="预览",
            form_name=form_name or schema.get("name") or "预览",
            form_desc=form_desc or (schema.get("descHTML") or schema.get("desc") or schema.get("description")),
            fields=fields_fallback,
            brand=brand
        )

# ========= 公共页：保存草稿（含文件）=========
@app.post("/site/<site_name>/draft/save")
def save_public_draft(site_name):
    conn = get_conn(); c = conn.cursor()

    # 读 schema & 上传上限
    c.execute("SELECT schema_json FROM form_defs WHERE site_name=%s", (site_name,))
    row = c.fetchone()
    if not row:
        conn.close()
        return jsonify(ok=False, error="no such site"), 404

    schema = row[0] if isinstance(row[0], dict) else json.loads(row[0] or "{}")
    upload_cfg = (schema.get("upload") or (schema.get("settings") or {}).get("upload") or {})
    max_files = int(upload_cfg.get("max_files") or 3)
    allowed = set(
        x.strip().lower()
        for x in str(upload_cfg.get("allowed_file_types", "")).split(",")
        if x.strip()
    )

    token = (request.form.get("__draft_token") or uuid4().hex)

    # 非文件字段
    data = request.form.to_dict()
    data.pop("__draft_token", None)

    # 已上传 URL（由前端隐藏域传回）
    uploaded_map = {}
    for k in list(request.form.keys()):
        if k.startswith("__uploaded__"):
            field = k[len("__uploaded__"):]
            uploaded_map[field] = request.form.getlist(k)
            data.pop(k, None)

    # 保存新选择文件 -> URL
    site_folder = os.path.join(app.config["UPLOAD_FOLDER"], site_name)
    os.makedirs(site_folder, exist_ok=True)

    files_payload = {**uploaded_map}  # field -> [urls]
    for field_key in request.files:
        urls = files_payload.get(field_key, [])[:]
        remain = max(0, max_files - len(urls))
        if remain <= 0:
            continue
        for f in request.files.getlist(field_key)[:remain]:
            if not f or not f.filename:
                continue
            ext = Path(f.filename).suffix.lower().lstrip(".")
            if allowed and ext not in allowed:
                continue
            uniq = f"{int(time())}_{uuid4().hex[:8]}_{secure_filename(f.filename)}"
            abs_path = os.path.join(site_folder, uniq)
            f.save(abs_path)
            urls.append(f"/site/{site_name}/uploads/{uniq}")
        files_payload[field_key] = urls

    # 建表 & UPSERT
    schema_name = _safe_schema(site_name)
    c.execute(f'''
        CREATE TABLE IF NOT EXISTS "{schema_name}".drafts (
          token TEXT PRIMARY KEY,
          data  JSONB,
          files JSONB,
          created_at TIMESTAMP DEFAULT NOW(),
          updated_at TIMESTAMP DEFAULT NOW()
        )
    ''')
    c.execute(
        f'''INSERT INTO "{schema_name}".drafts(token, data, files)
            VALUES (%s, %s, %s)
            ON CONFLICT (token) DO UPDATE
            SET data=EXCLUDED.data, files=EXCLUDED.files, updated_at=NOW()''',
        (token, json.dumps(data, ensure_ascii=False), json.dumps(files_payload, ensure_ascii=False))
    )
    conn.commit(); conn.close()

    return jsonify(ok=True, token=token, files=files_payload)

def _extract_email(data: dict) -> str:
    if not isinstance(data, dict): return ""
    for key in ("email", "邮箱", "mail"):
        v = data.get(key)
        if v: return str(v)
    return ""

@app.route("/site/<site_name>/admin/export_word/<int:sub_id>")
@admin_required
def export_word(site_name, sub_id):
    schema = _safe_schema(site_name)
    conn = get_conn(); c = conn.cursor()
    c.execute(f'SET search_path TO "{schema}", public')
    c.execute("SELECT data FROM submissions WHERE id=%s", (sub_id,))
    row = c.fetchone(); conn.close()
    if not row: return "❌ 记录不存在", 404

    data = row[0] if isinstance(row[0], dict) else (json.loads(row[0]) if row[0] else {})
    data = _normalize_obj(data)

    doc = Document(); doc.add_heading(f"提交 #{sub_id}", level=1)
    # 注意：此处遍历 schema["fields"] 可能导致 KeyError（若 schema 未带 fields）。保持原样不动。
    for k, v in (data or {}).items():
        safe_k = _maybe_fix_encoding(str(k))
        safe_v = _maybe_fix_encoding("" if v is None else str(v))
        p = doc.add_paragraph()
        p.add_run(f"{safe_k}: ").bold = True
        p.add_run(safe_v)

    buffer = io.BytesIO(); doc.save(buffer); buffer.seek(0)
    return send_file(
        buffer, as_attachment=True,
        download_name=f"submission_{sub_id}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.route("/site/<site_name>/admin/export_excel/<int:sub_id>")
@admin_required
def export_excel(site_name, sub_id):
    schema = _safe_schema(site_name)
    conn = get_conn(); c = conn.cursor()
    c.execute(f'SET search_path TO "{schema}", public')
    c.execute("SELECT data FROM submissions WHERE id=%s", (sub_id,))
    row = c.fetchone(); conn.close()
    if not row: return "❌ 记录不存在", 404
    data = row[0] if isinstance(row[0], dict) else (json.loads(row[0]) if row[0] else {})
    rows = [(_maybe_fix_encoding(str(k)),
             _maybe_fix_encoding("" if v is None else str(v)))
            for k, v in data.items()]
    df = pd.DataFrame(rows, columns=["字段", "内容"])

    try:
        buffer = io.BytesIO()
        with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
            df.to_excel(writer, index=False)
        buffer.seek(0)
        return send_file(
            buffer, as_attachment=True,
            download_name=f"submission_{sub_id}.xlsx",
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
    except Exception:
        csv_io = io.StringIO()
        df.to_csv(csv_io, index=False)
        mem = io.BytesIO(csv_io.getvalue().encode("utf-8-sig"))
        return send_file(
            mem, as_attachment=True,
            download_name=f"submission_{sub_id}.csv",
            mimetype="text/csv; charset=utf-8"
        )

def try_fix(s):
    if not isinstance(s, str):
        return s
    try:
        return s.encode("latin1").decode("utf-8")
    except Exception:
        return s

@app.route("/site/<site_name>/admin/api/export_all_excel")
@admin_required
def export_all_excel(site_name):
    schema = _safe_schema(site_name)
    conn = get_conn(); c = conn.cursor()
    c.execute(f'SET search_path TO "{schema}", public')
    c.execute("SELECT id, data, status, review_comment, created_at FROM submissions ORDER BY id")
    rows = c.fetchall(); conn.close()

    records, cols = [], {"id", "status", "review_comment", "created_at"}
    for r in rows:
        d = r[1] if isinstance(r[1], dict) else (json.loads(r[1]) if r[1] else {})
        d = _normalize_obj(d)
        rec = {
            "id": r[0],
            "status": _maybe_fix_encoding(r[2] or ""),
            "review_comment": _maybe_fix_encoding(r[3] or ""),
            "created_at": str(r[4]) if r[4] else ""
        }
        for k, v in (d or {}).items():
            kk = _maybe_fix_encoding(str(k))
            rec[kk] = _maybe_fix_encoding("" if v is None else str(v))
            cols.add(kk)
        records.append(rec)

    # 统一列顺序，保证固定列在前
    fixed = ["id", "status", "review_comment", "created_at"]
    dynamic = [col for col in cols if col not in fixed]
    df = pd.DataFrame(records, columns=fixed + list(dynamic))

    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        df.to_excel(writer, index=False)
    buf.seek(0)
    return send_file(
        buf, as_attachment=True,
        download_name=f"{site_name}_all_submissions.xlsx",
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

@app.route("/site/<site_name>/admin/api/gallery")
@admin_required
def api_gallery(site_name):
    schema = _safe_schema(site_name)
    conn = get_conn(); c = conn.cursor()
    c.execute(f'SET search_path TO "{schema}", public')
    c.execute("SELECT data FROM submissions ORDER BY id DESC LIMIT 1000")
    rows = c.fetchall(); conn.close()

    image_keys = ("image","img","photo","picture","图片","照片","附件","attachment")
    exts = (".png",".jpg",".jpeg",".gif",".webp",".bmp",".svg")
    def _to_url(u: str):
        if not u: return None
        if u.startswith("http"): return u
        if u.startswith("/"):    return u      # 已是站内绝对路径：/site/<site>/uploads/xxx
        # 其他情况（历史数据可能是纯文件名），走全局 /uploads/<filename>
        return url_for("uploaded_file", filename=u, _external=True)

    items = []
    for (d,) in rows:
        data = d if isinstance(d, dict) else (json.loads(d) if d else {})
        for k, v in (data or {}).items():
            key_l = str(k).lower()
            cand_urls = []
            if isinstance(v, str):
                cand_urls = [v]
            elif isinstance(v, list):
                cand_urls = [x for x in v if isinstance(x, str)]
            else:
                continue
            for u in cand_urls:
                if (any(x in key_l for x in image_keys)) or u.lower().endswith(exts):
                    url = _to_url(u)
                    if url:
                        items.append({"field": k, "url": url})
    return jsonify({"ok": True, "items": items})

def drop_bg_notify_from_all():
    conn = get_conn(); c = conn.cursor()
    c.execute("""
      UPDATE form_defs
         SET schema_json = jsonb_strip_nulls(
             (schema_json - 'bg' - 'bg_position' - 'notifications')
           )
    """)
    conn.commit(); conn.close()

def _norm(s: str) -> str:
    """把字符串做宽松匹配：去空白/标点并小写。"""
    import re as _re
    return _re.sub(r'[\s\-\_\.\|：:（）\(\)]+', '', str(s or '')).lower()

def _extract_label(f: dict) -> str:
    """从字段定义里取用户看到的标题/标签。"""
    if not isinstance(f, dict):
        return ''
    for k in ('label','title','text','name','placeholder','question','desc','description'):
        v = f.get(k)
        if isinstance(v, str) and v.strip():
            return v.strip()
    # 兼容常见嵌套
    for path in (('ui','label'),('ui','title'),('props','label'),('props','title'),
                 ('meta','label'),('meta','title')):
        cur = f
        for p in path:
            if not isinstance(cur, dict):
                cur = None; break
            cur = cur.get(p)
        if isinstance(cur, str) and cur.strip():
            return cur.strip()
    return ''

# === 变更点 ④：图表配置读写 + 按配置返回图表数据 ===

@app.route("/site/<site_name>/admin/api/charts_config", methods=["GET", "POST"])
@admin_required
def api_charts_config(site_name):
    """
    GET:  返回保存的图表配置 {"charts":[{"field":"字段key","type":"pie|line|flow","label":"可选显示名"}]}
    POST: 保存图表配置，body 形如 {"charts":[...]}；会写回 form_defs.schema_json.charts_config
    """
    conn = get_conn(); c = conn.cursor()
    try:
        if request.method == "GET":
            c.execute("SELECT schema_json FROM form_defs WHERE site_name=%s", (site_name,))
            row = c.fetchone()
            if not row:
                return jsonify({"ok": False, "error": "不存在的表单"}), 404
            schema = row[0] if isinstance(row[0], dict) else (json.loads(row[0]) if row[0] else {})
            cfg = schema.get("charts_config") or {"charts": []}
            return jsonify({"ok": True, "config": cfg})

        # POST 保存
        payload = request.get_json(silent=True) or {}
        charts = payload.get("charts") or []
        # 轻校验
        norm = []
        for ch in charts:
            if not isinstance(ch, dict):
                continue
            field = (ch.get("field") or "").strip()
            ctype = (ch.get("type") or "pie").lower()
            label = (ch.get("label") or "").strip()
            if not field:
                continue
            if ctype not in ("pie", "line", "flow"):
                ctype = "pie"
            norm.append({"field": field, "type": ctype, "label": label})
        c.execute("SELECT schema_json FROM form_defs WHERE site_name=%s", (site_name,))
        row = c.fetchone()
        if not row:
            return jsonify({"ok": False, "error": "不存在的表单"}), 404
        schema = row[0] if isinstance(row[0], dict) else (json.loads(row[0]) if row[0] else {})
        schema["charts_config"] = {"charts": norm}
        c.execute("UPDATE form_defs SET schema_json=%s WHERE site_name=%s", (Json(schema), site_name))
        conn.commit()
        return jsonify({"ok": True})
    except Exception as e:
        conn.rollback()
        return jsonify({"ok": False, "error": str(e)}), 500
    finally:
        conn.close()

@app.route("/site/<site_name>/admin/api/charts", methods=["GET"])
@admin_required
def api_charts(site_name):
    from collections import Counter, defaultdict
    from datetime import datetime, timedelta

    schema_name = _safe_schema(site_name)

    # 最近提交
    conn = get_conn(); c = conn.cursor()
    try:
        c.execute(f'SET search_path TO "{schema_name}", public')
        c.execute("""
            SELECT data, status, created_at
            FROM submissions
            ORDER BY id DESC
            LIMIT 2000
        """)
        rows = c.fetchall()
    except Exception as e:
        conn.rollback()
        return jsonify({"ok": False, "error": str(e)}), 500
    finally:
        conn.close()

    # 读 schema 与图表配置
    conn2 = get_conn(); c2 = conn2.cursor()
    try:
        c2.execute("SELECT schema_json FROM form_defs WHERE site_name=%s", (site_name,))
        r = c2.fetchone()
    finally:
        conn2.close()
    schema_json = r[0] if (r and isinstance(r[0], dict)) else (json.loads(r[0]) if r and r[0] else {})
    charts_cfg = (schema_json.get("charts_config") or {}).get("charts") or []

    # === 新增：构建「标题 -> 字段key」映射，支持用题目标题查图 ===
    fields = (schema_json or {}).get("fields", []) or []
    key_set = set()
    label_to_key = {}
    for f in fields:
        k = f.get("key") or f.get("id") or f.get("name")
        if not k:
            continue
        key_set.add(str(k))
        lab = _extract_label(f)
        if lab:
            label_to_key[_norm(lab)] = str(k)

    def resolve_field(user_input: str) -> str | None:
        """支持：直接传 key；或传中文标题/英文标题"""
        if not user_input:
            return None
        if user_input in key_set:
            return user_input
        return label_to_key.get(_norm(user_input))

    # 查询参数（单图临时查看）
    q_field_raw = (request.args.get("field") or "").strip()
    q_type  = (request.args.get("type") or "").strip().lower()
    if q_type not in ("pie","line","flow",""):
        q_type = "pie"
    q_field = resolve_field(q_field_raw) if q_field_raw else ""

    # 聚合容器
    def to_dict(obj):
        try:
            return obj if isinstance(obj, dict) else (json.loads(obj) if obj else {})
        except Exception:
            return {}

    now = datetime.utcnow()
    start_day = (now - timedelta(days=13)).date()
    daily = Counter()
    status_counter = Counter()
    field_counters = defaultdict(Counter)   # field_key -> Counter()

    # 扫描数据
    for data, status, created_at in rows:
        d = to_dict(data)
        try:
            dt = created_at if isinstance(created_at, datetime) else datetime.fromisoformat(str(created_at))
        except Exception:
            dt = now
        day = dt.date()
        if day >= start_day:
            daily[day.isoformat()] += 1

        s = (status or "").strip() or "待审核"
        status_counter[s] += 1

        for k, v in (d or {}).items():
            k = str(k)
            if isinstance(v, list):
                for each in v:
                    field_counters[k][str(each)] += 1
            else:
                field_counters[k][str(v)] += 1

    # 通用结果
    dates = [(now - timedelta(days=i)).date() for i in range(13, -1, -1)]
    daily_arr = [{"date": d.isoformat(), "count": int(daily.get(d.isoformat(), 0))} for d in dates]
    status_arr = [{"name": k, "count": int(v)} for k, v in status_counter.items()]

    def chart_payload(field_key: str, ctype: str, label: str = None):
        label = label or field_key or "字段"
        dist = field_counters.get(field_key, Counter())
        cat = [{"value": k, "count": int(v)} for k, v in dist.items()]
        payload = {"field": field_key, "label": label, "type": ctype or "pie", "data": cat}
        if ctype == "line":
            payload["daily"] = daily_arr
        if ctype == "flow":
            payload["funnel"] = sorted(cat, key=lambda x: x["count"], reverse=True)
        return payload

    # 先按已保存配置返回
    charts = []
    if charts_cfg and not q_field:
        for ch in charts_cfg:
            want = (ch.get("field") or ch.get("label") or "").strip()
            fld  = resolve_field(want)
            if not fld:
                continue
            # 没给 label 时，用 schema 里的标题
            show_label = (ch.get("label") or
                          _extract_label(next((f for f in fields if (f.get("key") or f.get("id") or f.get("name")) == fld), {})) or
                          fld)
            charts.append(chart_payload(fld, (ch.get("type") or "pie").lower(), show_label))

    # 临时查看（query 覆盖）
    if q_field:
        charts = [chart_payload(q_field, q_type or "pie",
                                _extract_label(next((f for f in fields if (f.get("key") or f.get("id") or f.get("name")) == q_field), {})) or q_field)]

    # 兼容：没配置时自动挑一个字段
    if not charts:
        field_key, field_label = None, None
        for f in fields:
            t = (f.get("type") or "").lower()
            if t in ("select", "radio", "checkbox"):
                field_key = f.get("key") or f.get("id") or f.get("name")
                field_label = _extract_label(f) or field_key
                break
        if not field_key and rows:
            sample = to_dict(rows[0][0])
            if isinstance(sample, dict) and sample:
                field_key = next(iter(sample.keys()), None)
                field_label = field_key or "字段"
        if field_key:
            charts = [chart_payload(str(field_key), "pie", str(field_label))]

    resp = {
        "ok": True,
        "daily": daily_arr,
        "status": status_arr,
    }
    if charts:
        resp["charts"] = charts
        resp["field"] = {"label": charts[0].get("label") or "字段", "data": charts[0].get("data", [])}
    else:
        resp["charts"] = []
        resp["field"] = {"label": "字段", "data": []}

    return jsonify(resp)


# === 变更点 ④ 结束 ===

@app.route("/site/<site_name>/admin/api/charts_old", methods=["GET"])
@admin_required
def api_charts_old(site_name):
    # 为安全保留一个兼容端点（如果你的前端没用就忽略）
    return api_charts(site_name)

@app.route("/site/<site_name>/create_success")
@admin_required
def create_success(site_name):
    public_url = url_for('public_form', site_name=site_name)
    admin_url  = url_for('create_form', site=site_name)
    return render_template("create_success.html",
                           site_name=site_name,
                           public_url=public_url,
                           admin_url=admin_url)

@app.after_request
def allow_embed(resp):
    if request.endpoint == 'create_success':
        try:
            resp.headers.pop('X-Frame-Options', None)
        except Exception:
            pass
        try:
            if resp.content_type and resp.content_type.startswith('text/html'):
                html = resp.get_data(as_text=True)
                low = html.lower()
                if '</head>' in low and '<base' not in low:
                    html = html.replace('</head>', '<base target="_top"></head>', 1)
                    resp.set_data(html)
                elif '<body' in low and '<base' not in low:
                    html = html.replace('<body', '<base target="_top"><body', 1)
                    resp.set_data(html)
        except Exception:
            pass
    return resp

def _has_cjk(text: str) -> bool:
    if not text: return False
    import re as _re
    return bool(_re.search(r"[\u4e00-\u9fff]", str(text)))

# ========== 健康检查 ==========
@app.route("/_health")
def _health(): return "ok", 200

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
